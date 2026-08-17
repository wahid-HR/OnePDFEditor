#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
One PDF Editor v1.0
Made by Yumdrop Tech. Studio – 2026
Offline PDF / Image / Document viewer & editor for Windows.
"""

import os
import sys
import io
import tempfile
from pathlib import Path
from typing import Optional, List, Tuple, Dict, Any
from datetime import datetime

try:
    import pymupdf as fitz
except ImportError:
    import fitz  # type: ignore

from PIL import Image, ImageDraw, ImageTk, ImageEnhance, ImageFilter
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog, colorchooser

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

APP_NAME = "One PDF Editor"
APP_VERSION = "1.0"
APP_YEAR = "2026"
APP_STUDIO = "Yumdrop Tech. Studio"


def resource_path(*parts):
    """Resolve path for source, onedir, and onefile PyInstaller builds."""
    candidates = []
    if getattr(sys, "frozen", False):
        # onefile / onedir internal
        if hasattr(sys, "_MEIPASS"):
            candidates.append(Path(sys._MEIPASS))
        # onedir: next to the .exe
        candidates.append(Path(sys.executable).resolve().parent)
        # sometimes assets sit under _internal next to exe
        candidates.append(Path(sys.executable).resolve().parent / "_internal")
    else:
        candidates.append(Path(__file__).resolve().parent.parent)
        candidates.append(Path(__file__).resolve().parent)
    for base in candidates:
        p = base.joinpath(*parts)
        if p.exists():
            return p
    # fallback last candidate
    return candidates[0].joinpath(*parts)

MAX_UNDO = 12
MIN_FONT_SIZE = 4.0
DEFAULT_ZOOM = 1.0

BASE_FONTS = {
    "helv": "helv", "helvetica": "helv", "arial": "helv", "sans": "helv",
    "times": "times", "times-roman": "times", "serif": "times",
    "cour": "cour", "courier": "cour", "mono": "cour",
}

COLORS = {
    "bg": "#1e1e2e",
    "surface": "#2a2a3c",
    "surface2": "#34344a",
    "accent": "#7c5cff",
    "accent_hover": "#9b82ff",
    "text": "#e8e8f0",
    "text_dim": "#a0a0b8",
    "canvas_bg": "#12121a",
    "toolbar": "#252536",
}


def safe_color_to_rgb(color_int):
    if color_int is None:
        return (0.0, 0.0, 0.0)
    r = ((color_int >> 16) & 0xFF) / 255.0
    g = ((color_int >> 8) & 0xFF) / 255.0
    b = (color_int & 0xFF) / 255.0
    return (r, g, b)


def flags_to_style(flags):
    return {
        "italic": bool(flags & 2**1),
        "bold": bool(flags & 2**4),
        "mono": bool(flags & 2**3),
        "serif": bool(flags & 2**2),
    }


def choose_fallback_font(original_name, flags, text=""):
    name = (original_name or "").lower()
    style = flags_to_style(flags)
    if "+" in name:
        name = name.split("+", 1)[-1]
    for key, alias in BASE_FONTS.items():
        if key in name:
            return alias
    if style["mono"]:
        return "cour"
    if style["serif"] or "times" in name or "roman" in name:
        return "times"
    return "helv"


def text_needs_complex_script(text):
    for ch in text:
        o = ord(ch)
        if o > 0x00FF and not (0x2000 <= o <= 0x206F) and not (0x2010 <= o <= 0x2027):
            return True
    return False


def text_has_bengali(text):
    """True if text contains Bengali Unicode letters."""
    for ch in text:
        o = ord(ch)
        if 0x0980 <= o <= 0x09FF:  # Bengali block
            return True
        if 0x200C <= o <= 0x200D:  # ZWNJ/ZWJ often used in Bangla
            continue
    return False


# Minimal Bijoy (ANSI) → Unicode map for common Bangla typing
# Users typing with Bijoy keyboard layout on Windows often produce these codes.
_BIJOY_MAP = {
    "A": "অ", "i": "ই", "I": "ঈ", "u": "উ", "U": "ঊ", "e": "এ", "E": "ঐ",
    "o": "ও", "O": "ঔ", "k": "ক", "K": "খ", "g": "গ", "G": "ঘ", "c": "চ",
    "C": "ছ", "j": "জ", "J": "ঝ", "T": "ট", "t": "ত", "d": "দ", "D": "ড",
    "n": "ন", "N": "ণ", "p": "প", "P": "ফ", "f": "ফ", "b": "ব", "v": "ভ",
    "m": "ম", "z": "য", "r": "র", "l": "ল", "S": "শ", "s": "স", "h": "হ",
    "R": "ড়", "y": "য়", "w": "ৎ", "Y": "য়",
    "a": "া", "w": "্", "x": "ঁ", "X": "ঁ",
    "1": "১", "2": "২", "3": "৩", "4": "৪", "5": "৫",
    "6": "৬", "7": "৭", "8": "৮", "9": "৯", "0": "০",
}


def bijoy_to_unicode(text):
    """Best-effort Bijoy→Unicode. If already Unicode Bangla, return as-is."""
    if text_has_bengali(text):
        return text
    out = []
    for ch in text:
        out.append(_BIJOY_MAP.get(ch, ch))
    return "".join(out)


def get_bengali_font_path(bold=False):
    name = "NotoSansBengali-Bold.ttf" if bold else "NotoSansBengali-Regular.ttf"
    p = resource_path("assets", "fonts", name)
    if p.exists():
        return str(p)
    # try sibling
    p2 = resource_path("fonts", name)
    if p2.exists():
        return str(p2)
    return None


def measure_text_width(text, fontname, fontsize):
    try:
        font = fitz.Font(fontname)
        return font.text_length(text, fontsize=fontsize)
    except Exception:
        return len(text) * fontsize * 0.5


class PDFDocument:
    def __init__(self):
        self.doc = None
        self.path = None
        self.dirty = False
        self.source_type = "pdf"
        self._undo_stack = []
        self._redo_stack = []

    def open(self, path, password=""):
        ext = Path(path).suffix.lower()
        self.close()
        try:
            if ext in (".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tiff", ".tif"):
                return self._open_image(path)
            if ext == ".docx" and HAS_DOCX:
                return self._open_docx(path)
            self.doc = fitz.open(path)
            if self.doc.is_encrypted:
                if not self.doc.authenticate(password):
                    self.doc.close()
                    self.doc = None
                    return False
            self.path = path
            self.source_type = "pdf"
            self.dirty = False
            self._undo_stack.clear()
            self._push_undo()
            return True
        except Exception as e:
            raise RuntimeError(f"Cannot open file: {e}") from e

    def _open_image(self, path):
        img = Image.open(path)
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        # Offline quality enhance
        try:
            img = ImageEnhance.Sharpness(img).enhance(1.35)
            img = ImageEnhance.Contrast(img).enhance(1.12)
            img = ImageEnhance.Color(img).enhance(1.05)
        except Exception:
            pass
        bio = io.BytesIO()
        img.save(bio, format="PDF", resolution=150.0)
        bio.seek(0)
        self.doc = fitz.open("pdf", bio.read())
        # Best-effort OCR text layer (needs Tesseract on system)
        try:
            page = self.doc[0]
            tp = page.get_textpage_ocr(dpi=150, full=True)
            # OCR text becomes selectable/editable spans
            _ = page.get_text("dict", textpage=tp)
        except Exception:
            pass  # no tesseract — image-only PDF still works
        self.path = path
        self.source_type = "image"
        self.dirty = False
        self._undo_stack.clear()
        self._push_undo()
        return True

    def _open_docx(self, path):
        if not HAS_DOCX:
            raise RuntimeError("Word support requires python-docx")
        document = DocxDocument(path)
        pdf = fitz.open()
        page = pdf.new_page(width=595, height=842)
        y = 50
        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                y += 12
                continue
            while text and y < 800:
                chunk = text[:90]
                text = text[90:]
                page.insert_text((50, y), chunk, fontsize=11, fontname="helv")
                y += 16
            if y >= 800:
                page = pdf.new_page(width=595, height=842)
                y = 50
        self.doc = pdf
        self.path = path
        self.source_type = "docx"
        self.dirty = True
        self._undo_stack.clear()
        self._push_undo()
        return True

    def close(self):
        if self.doc:
            try:
                self.doc.close()
            except Exception:
                pass
        self.doc = None
        self.path = None
        self.dirty = False
        self.source_type = "pdf"
        self._undo_stack.clear()
        self._redo_stack.clear()

    def page_count(self):
        return len(self.doc) if self.doc else 0

    def get_page(self, idx):
        if self.doc and 0 <= idx < len(self.doc):
            return self.doc[idx]
        return None

    def _serialize(self):
        if not self.doc:
            return b""
        return self.doc.tobytes(garbage=3, deflate=True)

    def _push_undo(self):
        data = self._serialize()
        if data:
            self._undo_stack.append(data)
            if len(self._undo_stack) > MAX_UNDO:
                self._undo_stack.pop(0)
            self._redo_stack.clear()

    def save_state(self):
        self._push_undo()
        self.dirty = True

    def undo(self):
        if len(self._undo_stack) < 2 or not self.doc:
            return False
        current = self._undo_stack.pop()
        self._redo_stack.append(current)
        self._restore(self._undo_stack[-1])
        self.dirty = True
        return True

    def redo(self):
        if not self._redo_stack or not self.doc:
            return False
        data = self._redo_stack.pop()
        self._undo_stack.append(data)
        self._restore(data)
        self.dirty = True
        return True

    def _restore(self, data):
        if not data:
            return
        try:
            newdoc = fitz.open("pdf", data)
            if self.doc:
                self.doc.close()
            self.doc = newdoc
        except Exception:
            pass

    def save(self, path=None):
        if not self.doc:
            return False
        target = path or self.path
        if not target:
            return False
        if not target.lower().endswith(".pdf"):
            target = str(Path(target).with_suffix(".pdf"))
        try:
            self.doc.save(target, garbage=3, deflate=True, incremental=False)
            self.path = target
            self.dirty = False
            self.source_type = "pdf"
            return True
        except Exception as e:
            raise RuntimeError(f"Save failed: {e}") from e

    def get_text_spans(self, page_idx):
        page = self.get_page(page_idx)
        if not page:
            return []
        spans = []
        try:
            blocks = page.get_text("dict", flags=fitz.TEXTFLAGS_TEXT)["blocks"]
            for b in blocks:
                if b.get("type") != 0:
                    continue
                for line in b.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip("\x00")
                        if not text.strip():
                            continue
                        spans.append({
                            "text": text,
                            "bbox": fitz.Rect(span["bbox"]),
                            "origin": fitz.Point(span.get("origin", span["bbox"][:2])),
                            "font": span.get("font", "helv"),
                            "size": float(span.get("size", 11)),
                            "color": span.get("color", 0),
                            "flags": span.get("flags", 0),
                        })
        except Exception:
            pass
        return spans

    def search(self, query):
        results = []
        if not self.doc or not query:
            return results
        for i in range(len(self.doc)):
            try:
                for r in self.doc[i].search_for(query, quads=False):
                    results.append((i, fitz.Rect(r)))
            except Exception:
                continue
        return results

    def replace_span(self, page_idx, span, new_text, fontsize=None, bold=False, italic=False):
        page = self.get_page(page_idx)
        if not page or not new_text:
            return False, "Invalid"
        self.save_state()
        bbox = fitz.Rect(span["bbox"])
        origin = fitz.Point(span["origin"])
        if fontsize is None:
            fontsize = float(span["size"])
        color = safe_color_to_rgb(span["color"])
        flags = span.get("flags", 0)
        if bold:
            flags |= 2**4
        if italic:
            flags |= 2**1
        fontname = span.get("font", "helv")
        use_font = choose_fallback_font(fontname, flags, new_text)
        # Prefer styled base fonts when bold/italic requested
        if bold and italic:
            if use_font == "helv":
                use_font = "heit"  # Helvetica-Oblique approx; fall back
            # Base14 has limited styled names; use morph or just size
        complex_script = text_needs_complex_script(new_text)
        try:
            _ = fitz.Font(use_font)
        except Exception:
            use_font = "helv"
        target_width = max(bbox.width, 1.0)
        fitted_size = float(fontsize)
        try:
            width = measure_text_width(new_text, use_font, fitted_size)
            if width > target_width * 1.15:
                fitted_size = max(MIN_FONT_SIZE, fitted_size * (target_width / width) * 0.98)
        except Exception:
            pass
        # Sample approximate background from page render near the span
        bg_fill = (1, 1, 1)
        try:
            pix = page.get_pixmap(clip=bbox, matrix=fitz.Matrix(2, 2), alpha=False)
            if pix.width > 2 and pix.height > 2 and pix.samples:
                # average a few corner pixels (avoid center which is text)
                samples = []
                w, h = pix.width, pix.height
                for (px, py) in [(1, 1), (w-2, 1), (1, h-2), (w-2, h-2), (w//2, 1), (w//2, h-2)]:
                    i = (py * w + px) * 3
                    if i + 2 < len(pix.samples):
                        samples.append((pix.samples[i]/255.0, pix.samples[i+1]/255.0, pix.samples[i+2]/255.0))
                if samples:
                    bg_fill = (
                        sum(s[0] for s in samples) / len(samples),
                        sum(s[1] for s in samples) / len(samples),
                        sum(s[2] for s in samples) / len(samples),
                    )
        except Exception:
            pass
        try:
            page.add_redact_annot(bbox, fill=bg_fill)
            page.apply_redactions(
                images=fitz.PDF_REDACT_IMAGE_NONE,
                graphics=fitz.PDF_REDACT_LINE_ART_NONE,
                text=fitz.PDF_REDACT_TEXT_REMOVE,
            )
        except Exception as e:
            return False, f"Redaction failed: {e}"
        # Convert Bijoy-style input to Unicode if needed
        new_text = bijoy_to_unicode(new_text)
        is_bn = text_has_bengali(new_text) or text_needs_complex_script(new_text)
        fontfile = get_bengali_font_path(bold=bold) if is_bn else None

        try:
            if fontfile:
                # Register & use embedded Unicode font (Bangla etc.)
                try:
                    page.insert_font(fontname="bnfont", fontfile=fontfile)
                    fname = "bnfont"
                except Exception:
                    fname = use_font
                page.insert_text(
                    origin, new_text, fontname=fname, fontfile=fontfile,
                    fontsize=fitted_size, color=color, render_mode=0, overlay=True,
                )
                if bold:
                    page.insert_text(
                        fitz.Point(origin.x + 0.4, origin.y), new_text,
                        fontname=fname, fontfile=fontfile,
                        fontsize=fitted_size, color=color, render_mode=0, overlay=True,
                    )
            else:
                page.insert_text(
                    origin, new_text, fontname=use_font, fontsize=fitted_size,
                    color=color, render_mode=0, overlay=True,
                )
                if bold:
                    page.insert_text(
                        fitz.Point(origin.x + 0.35, origin.y), new_text,
                        fontname=use_font, fontsize=fitted_size,
                        color=color, render_mode=0, overlay=True,
                    )
        except Exception as e:
            # Last resort: textbox with fontfile
            try:
                if fontfile:
                    page.insert_textbox(
                        bbox, new_text, fontname="helv", fontfile=fontfile,
                        fontsize=fitted_size, color=color, align=0, overlay=True,
                    )
                else:
                    return False, f"Insert failed: {e}"
            except Exception as e2:
                return False, f"Insert failed: {e2}"
        self.dirty = True
        msg = f"Edited (size={fitted_size:.1f}"
        if bold:
            msg += ", bold"
        if is_bn:
            msg += ", Bangla"
        msg += ")"
        return True, msg


    def extract_page_images(self, page_idx):
        """Return list of (xref, PIL.Image) for images on the page."""
        page = self.get_page(page_idx)
        if not page or not self.doc:
            return []
        result = []
        try:
            for info in page.get_images(full=True):
                xref = info[0]
                try:
                    pix = fitz.Pixmap(self.doc, xref)
                    if pix.n - pix.alpha > 3:  # CMYK etc
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    mode = "RGBA" if pix.alpha else "RGB"
                    img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
                    result.append((xref, img))
                except Exception:
                    continue
        except Exception:
            pass
        return result

    def sample_rect_background(self, page_idx, rect):
        """Average background color from corners of a page region (0-1 RGB)."""
        page = self.get_page(page_idx)
        if not page:
            return (1.0, 1.0, 1.0)
        try:
            # slightly inset to avoid borders
            r = fitz.Rect(rect)
            if r.width < 2 or r.height < 2:
                return (1.0, 1.0, 1.0)
            pix = page.get_pixmap(clip=r, matrix=fitz.Matrix(2, 2), alpha=False)
            w, h = pix.width, pix.height
            if w < 2 or h < 2 or not pix.samples:
                return (1.0, 1.0, 1.0)
            pts = [(1, 1), (w - 2, 1), (1, h - 2), (w - 2, h - 2),
                   (w // 2, 1), (w // 2, h - 2), (1, h // 2), (w - 2, h // 2)]
            samples = []
            for px, py in pts:
                i = (py * w + px) * 3
                if i + 2 < len(pix.samples):
                    samples.append((
                        pix.samples[i] / 255.0,
                        pix.samples[i + 1] / 255.0,
                        pix.samples[i + 2] / 255.0,
                    ))
            if not samples:
                return (1.0, 1.0, 1.0)
            return (
                sum(s[0] for s in samples) / len(samples),
                sum(s[1] for s in samples) / len(samples),
                sum(s[2] for s in samples) / len(samples),
            )
        except Exception:
            return (1.0, 1.0, 1.0)

    def fill_rectangle(self, page_idx, rect, color_rgb=None):
        """Cover area to hide text. If color_rgb is None, auto-sample PDF background."""
        page = self.get_page(page_idx)
        if not page:
            return False
        if color_rgb is None:
            color_rgb = self.sample_rect_background(page_idx, rect)
        self.save_state()
        try:
            page.add_redact_annot(rect, fill=color_rgb)
            page.apply_redactions(
                images=fitz.PDF_REDACT_IMAGE_NONE,
                graphics=fitz.PDF_REDACT_LINE_ART_NONE,
                text=fitz.PDF_REDACT_TEXT_REMOVE,
            )
            shape = page.new_shape()
            shape.draw_rect(rect)
            shape.finish(color=color_rgb, fill=color_rgb, width=0)
            shape.commit()
            self.dirty = True
            return True
        except Exception:
            return False

    def insert_text_in_rect(self, page_idx, rect, text, fontsize=None, color=(0, 0, 0)):
        page = self.get_page(page_idx)
        if not page or not text:
            return False
        self.save_state()
        try:
            text = bijoy_to_unicode(text)
            r = fitz.Rect(rect)
            if fontsize is None:
                fontsize = max(6.0, min(24.0, r.height * 0.65))
            origin = fitz.Point(r.x0 + 1, r.y0 + fontsize * 0.85)
            fontfile = get_bengali_font_path() if text_has_bengali(text) else None
            kwargs = dict(fontsize=fontsize, color=color, render_mode=0, overlay=True)
            if fontfile:
                try:
                    page.insert_font(fontname="bnfont", fontfile=fontfile)
                except Exception:
                    pass
                page.insert_text(origin, text, fontname="bnfont", fontfile=fontfile, **kwargs)
            else:
                page.insert_text(origin, text, fontname="helv", **kwargs)
            self.dirty = True
            return True
        except Exception:
            return False

    def insert_signature_image(self, page_idx, img, rect):
        page = self.get_page(page_idx)
        if not page:
            return False
        self.save_state()
        try:
            if img.mode != "RGBA":
                img = img.convert("RGBA")
            bio = io.BytesIO()
            img.save(bio, format="PNG")
            bio.seek(0)
            page.insert_image(rect, stream=bio.getvalue(), overlay=True)
            self.dirty = True
            return True
        except Exception:
            return False


class SignatureDialog(tk.Toplevel):
    def __init__(self, parent, on_done):
        super().__init__(parent)
        self.title("Draw Signature")
        self.configure(bg=COLORS["surface"])
        self.resizable(True, True)
        self.transient(parent)
        self.grab_set()
        self.on_done = on_done
        self.strokes = []
        self.current_stroke = []
        self.drawing = False
        self.canvas = tk.Canvas(self, bg="white", width=520, height=220, cursor="pencil",
                                highlightthickness=1, highlightbackground=COLORS["accent"])
        self.canvas.pack(fill=tk.BOTH, expand=True, padx=12, pady=12)
        self.canvas.bind("<ButtonPress-1>", self._start)
        self.canvas.bind("<B1-Motion>", self._move)
        self.canvas.bind("<ButtonRelease-1>", self._end)
        btn_frame = tk.Frame(self, bg=COLORS["surface"])
        btn_frame.pack(fill=tk.X, padx=12, pady=(0, 12))
        self._btn(btn_frame, "Clear", self._clear).pack(side=tk.LEFT, padx=4)
        self._btn(btn_frame, "Cancel", self._cancel).pack(side=tk.RIGHT, padx=4)
        self._btn(btn_frame, "Use Signature", self._accept, accent=True).pack(side=tk.RIGHT, padx=4)
        self.protocol("WM_DELETE_WINDOW", self._cancel)
        self.geometry("560x320")
        self.focus_set()

    def _btn(self, parent, text, cmd, accent=False):
        bg = COLORS["accent"] if accent else COLORS["surface2"]
        return tk.Button(parent, text=text, command=cmd, bg=bg, fg=COLORS["text"],
                         activebackground=COLORS["accent_hover"], activeforeground="white",
                         relief=tk.FLAT, padx=14, pady=6, cursor="hand2", font=("Segoe UI", 10))

    def _start(self, event):
        self.drawing = True
        self.current_stroke = [(event.x, event.y)]
        self.canvas.create_oval(event.x-1, event.y-1, event.x+1, event.y+1, fill="black", outline="black")

    def _move(self, event):
        if not self.drawing:
            return
        x, y = event.x, event.y
        if self.current_stroke:
            px, py = self.current_stroke[-1]
            self.canvas.create_line(px, py, x, y, fill="black", width=2, capstyle=tk.ROUND, smooth=True)
        self.current_stroke.append((x, y))

    def _end(self, event):
        if self.drawing and self.current_stroke:
            self.strokes.append(self.current_stroke)
        self.drawing = False
        self.current_stroke = []

    def _clear(self):
        self.canvas.delete("all")
        self.strokes.clear()
        self.current_stroke = []

    def _cancel(self):
        self.on_done(None)
        self.destroy()

    def _accept(self):
        if not self.strokes:
            messagebox.showwarning("Empty", "Please draw a signature first.", parent=self)
            return
        w = int(self.canvas.winfo_width()) or 520
        h = int(self.canvas.winfo_height()) or 220
        img = Image.new("RGBA", (w, h), (0, 0, 0, 0))
        draw = ImageDraw.Draw(img)
        for stroke in self.strokes:
            if len(stroke) < 2:
                if stroke:
                    x, y = stroke[0]
                    draw.ellipse([x-1, y-1, x+1, y+1], fill=(0, 0, 0, 255))
                continue
            draw.line(stroke, fill=(0, 0, 0, 255), width=2, joint="curve")
        self.on_done(img)
        self.destroy()


class ScreenshotPopup(tk.Toplevel):
    def __init__(self, parent, image):
        super().__init__(parent)
        self.title("Screenshot")
        self.configure(bg=COLORS["surface"])
        self.transient(parent)
        self.grab_set()
        self.image = image
        preview = image.copy()
        preview.thumbnail((480, 360))
        self.photo = ImageTk.PhotoImage(preview)
        lbl = tk.Label(self, image=self.photo, bg=COLORS["surface"])
        lbl.pack(padx=16, pady=16)
        btn_frame = tk.Frame(self, bg=COLORS["surface"])
        btn_frame.pack(pady=(0, 16))
        tk.Button(btn_frame, text="Copy to Clipboard", command=self._copy,
                  bg=COLORS["accent"], fg="white", relief=tk.FLAT, padx=16, pady=8,
                  font=("Segoe UI", 10), cursor="hand2").pack(side=tk.LEFT, padx=8)
        tk.Button(btn_frame, text="Save as PNG", command=self._save,
                  bg=COLORS["surface2"], fg=COLORS["text"], relief=tk.FLAT, padx=16, pady=8,
                  font=("Segoe UI", 10), cursor="hand2").pack(side=tk.LEFT, padx=8)
        tk.Button(btn_frame, text="Close", command=self.destroy,
                  bg=COLORS["surface2"], fg=COLORS["text_dim"], relief=tk.FLAT, padx=12, pady=8,
                  font=("Segoe UI", 10), cursor="hand2").pack(side=tk.LEFT, padx=8)
        self.geometry("520x460")
        self.focus_set()

    def _copy(self):
        try:
            output = io.BytesIO()
            self.image.convert("RGB").save(output, "BMP")
            data = output.getvalue()[14:]
            output.close()
            try:
                import win32clipboard
                win32clipboard.OpenClipboard()
                win32clipboard.EmptyClipboard()
                win32clipboard.SetClipboardData(win32clipboard.CF_DIB, data)
                win32clipboard.CloseClipboard()
                messagebox.showinfo("Copied", "Screenshot copied to clipboard.", parent=self)
            except ImportError:
                tmp = Path(tempfile.gettempdir()) / "one_pdf_screenshot.png"
                self.image.save(tmp)
                messagebox.showinfo("Saved", f"Clipboard helper not available.\nSaved to:\n{tmp}", parent=self)
        except Exception as e:
            messagebox.showerror("Error", str(e), parent=self)

    def _save(self):
        path = filedialog.asksaveasfilename(
            parent=self, defaultextension=".png",
            filetypes=[("PNG image", "*.png"), ("JPEG", "*.jpg")],
            initialfile=f"screenshot_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png",
        )
        if path:
            self.image.save(path)
            messagebox.showinfo("Saved", f"Saved to:\n{path}", parent=self)
            self.destroy()



class BanglaKeyboard(tk.Toplevel):
    """On-screen Bangla Unicode keyboard — no system IME required."""

    # Rows of Bangla characters for a compact keyboard
    ROWS = [
        list("১২৩৪৫৬৭৮৯০"),
        list("অআইঈউঊঋএঐওঔ"),
        list("কখগঘঙচছজঝঞ"),
        list("টঠডঢণতথদধন"),
        list("পফবভমযরল"),
        list("শষসহড়ঢ়য়ৎ"),
        list("ািীুূৃেৈোৌ্ংঃঁ"),
    ]
    EXTRA = [
        ("ক্ষ", "ক্ষ"), ("জ্ঞ", "জ্ঞ"), ("ঞ্জ", "ঞ্জ"), ("ত্ত", "ত্ত"),
        ("Space", " "), ("⌫", "BACK"), ("Enter", "ENTER"),
    ]

    def __init__(self, parent, on_char, on_enter=None):
        super().__init__(parent)
        self.title("বাংলা কিবোর্ড")
        self.configure(bg=COLORS["surface"])
        self.transient(parent)
        self.on_char = on_char
        self.on_enter = on_enter
        self.resizable(False, False)

        tk.Label(
            self, text="ক্লিক করে বাংলা লিখুন (Avro লাগবে না)",
            bg=COLORS["surface"], fg=COLORS["text_dim"], font=("Segoe UI", 9),
        ).pack(pady=(8, 4))

        # Phonetic quick type
        pf = tk.Frame(self, bg=COLORS["surface"])
        pf.pack(fill=tk.X, padx=10, pady=4)
        tk.Label(pf, text="Phonetic:", bg=COLORS["surface"], fg=COLORS["text"],
                 font=("Segoe UI", 9)).pack(side=tk.LEFT)
        self.phon_var = tk.StringVar()
        pe = tk.Entry(pf, textvariable=self.phon_var, width=22,
                      bg=COLORS["surface2"], fg=COLORS["text"],
                      insertbackground=COLORS["text"], relief=tk.FLAT, font=("Segoe UI", 11))
        pe.pack(side=tk.LEFT, padx=4)
        pe.bind("<Return>", self._phonetic_apply)
        tk.Button(pf, text="→ বাংলা", command=self._phonetic_apply,
                  bg=COLORS["accent"], fg="white", relief=tk.FLAT, padx=8, pady=2,
                  font=("Segoe UI", 9), cursor="hand2").pack(side=tk.LEFT, padx=2)

        body = tk.Frame(self, bg=COLORS["surface"])
        body.pack(padx=8, pady=6)
        for row in self.ROWS:
            rf = tk.Frame(body, bg=COLORS["surface"])
            rf.pack(pady=1)
            for ch in row:
                self._key(rf, ch, ch)

        ef = tk.Frame(body, bg=COLORS["surface"])
        ef.pack(pady=4)
        for label, val in self.EXTRA:
            self._key(ef, label, val, wide=(val in (" ", "BACK", "ENTER")))

        tk.Button(self, text="Close", command=self.destroy,
                  bg=COLORS["surface2"], fg=COLORS["text_dim"], relief=tk.FLAT,
                  padx=12, pady=4, font=("Segoe UI", 9)).pack(pady=(0, 10))

        self.geometry("+%d+%d" % (parent.winfo_rootx() + 80, parent.winfo_rooty() + 120))

    def _key(self, parent, label, value, wide=False):
        w = 6 if wide else 3
        btn = tk.Button(
            parent, text=label, width=w,
            command=lambda v=value: self._press(v),
            bg=COLORS["surface2"], fg=COLORS["text"],
            activebackground=COLORS["accent"], activeforeground="white",
            relief=tk.FLAT, font=("Segoe UI", 11), cursor="hand2", padx=2, pady=2,
        )
        btn.pack(side=tk.LEFT, padx=1, pady=1)

    def _press(self, value):
        if value == "ENTER":
            if self.on_enter:
                self.on_enter()
            return
        self.on_char(value)

    def _phonetic_apply(self, event=None):
        raw = self.phon_var.get().strip()
        if not raw:
            return
        converted = phonetic_bangla(raw)
        self.on_char(converted)
        self.phon_var.set("")


# Lightweight English-phonetic → Bangla (common patterns, longest match first)
_PHONETIC_RULES = [
    ("ksh", "ক্ষ"), ("gg", "জ্ঞ"), ("ng", "ং"), ("nj", "ঞ্জ"), ("tt", "ত্ত"),
    ("th", "থ"), ("Th", "ঠ"), ("dh", "ধ"), ("Dh", "ঢ"), ("ch", "চ"), ("Ch", "ছ"),
    ("sh", "শ"), ("Sh", "ষ"), ("ph", "ফ"), ("bh", "ভ"), ("jh", "ঝ"),
    ("kh", "খ"), ("gh", "ঘ"), ("rh", "ঢ়"),
    ("ou", "ৌ"), ("oi", "ৈ"), ("ee", "ী"), ("oo", "ূ"),
    ("aa", "া"), ("ri", "ৃ"),
    ("a", "া"), ("i", "ি"), ("I", "ী"), ("u", "ু"), ("U", "ূ"),
    ("e", "ে"), ("o", "ো"), ("O", "ৌ"),
    ("k", "ক"), ("g", "গ"), ("c", "চ"), ("j", "জ"), ("t", "ত"), ("T", "ট"),
    ("d", "দ"), ("D", "ড"), ("n", "ন"), ("N", "ণ"), ("p", "প"), ("b", "ব"),
    ("m", "ম"), ("y", "য়"), ("r", "র"), ("l", "ল"), ("s", "স"), ("h", "হ"),
    ("z", "য"), ("f", "ফ"), ("v", "ভ"), ("w", "ও"),
    ("0", "০"), ("1", "১"), ("2", "২"), ("3", "৩"), ("4", "৪"),
    ("5", "৫"), ("6", "৬"), ("7", "৭"), ("8", "৮"), ("9", "৯"),
]


def phonetic_bangla(text):
    """Convert simple phonetic English to Bangla Unicode."""
    text = text.lower()
    out = []
    i = 0
    # Independent vowels at word start-ish
    indep = {
        "a": "অ", "aa": "আ", "i": "ই", "ii": "ঈ", "u": "উ", "uu": "ঊ",
        "e": "এ", "oi": "ঐ", "o": "ও", "ou": "ঔ", "ri": "ঋ",
    }
    while i < len(text):
        if text[i] in " \t\n":
            out.append(text[i])
            i += 1
            continue
        matched = False
        # try independent vowel if at start or after space
        if i == 0 or (i > 0 and text[i - 1] in " \t\n"):
            for ln in (2, 1):
                chunk = text[i:i + ln]
                if chunk in indep:
                    out.append(indep[chunk])
                    i += ln
                    matched = True
                    break
        if matched:
            continue
        for lat, bn in _PHONETIC_RULES:
            if text.startswith(lat, i):
                out.append(bn)
                i += len(lat)
                matched = True
                break
        if not matched:
            out.append(text[i])
            i += 1
    return "".join(out)


class OnePDFEditor(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(f"{APP_NAME}  v{APP_VERSION}")
        self.geometry("1200x800")
        self.minsize(900, 600)
        self.configure(bg=COLORS["bg"])
        self._set_app_icon()
        self.pdf = PDFDocument()
        self.current_page = 0
        self.zoom = DEFAULT_ZOOM
        self.photo = None
        self.search_results = []
        self.search_index = -1
        self.highlight_rect = None
        self.selected_span = None
        self.selected_page = -1
        self.signature_img = None
        self.sig_rect = None
        self.placing_signature = False
        self._drag_start = None
        self._edit_entry = None
        self._edit_span = None
        # Color fill box tool
        self.fill_mode = False
        self.pick_color_mode = False
        self.copy_sign_mode = False
        self.screenshot_mode = False
        self.screenshot_rect = None
        self.fill_color = None  # None = auto-sample PDF background
        self.fill_color_manual = False  # True only after Pick Color
        self.text_color = (0.0, 0.0, 0.0)  # RGB 0-1 for new text
        self.fill_rect = None
        self.copy_sign_rect = None
        self.copied_signatures = []
        self._dash_photos = []
        self._dash_idx = 0
        self._dash_job = None
        self._dash_bg_photo = None
        self._build_ui()
        self.after(100, self.render_page)
        self._bind_shortcuts()
        self.protocol("WM_DELETE_WINDOW", self._on_close)
        self.update_idletasks()
        x = (self.winfo_screenwidth() - self.winfo_width()) // 2
        y = (self.winfo_screenheight() - self.winfo_height()) // 2
        self.geometry(f"+{x}+{y}")

    def _set_app_icon(self):
        """Set window/taskbar icon from packaged assets (title bar & taskbar only)."""
        try:
            png = resource_path("assets", "icon_256.png")
            if not png.exists():
                png = resource_path("assets", "icon_64.png")
            if png.exists():
                img = Image.open(png)
                self._app_icon_img = ImageTk.PhotoImage(img)
                self.iconphoto(True, self._app_icon_img)
            ico = resource_path("assets", "OnePDFEditor.ico")
            if ico.exists() and sys.platform.startswith("win"):
                try:
                    self.iconbitmap(default=str(ico))
                except Exception:
                    try:
                        self.iconbitmap(str(ico))
                    except Exception:
                        pass
        except Exception:
            pass

    def _make_tool_btn(self, parent, text, command):
        return tk.Button(
            parent, text=text, command=command,
            bg=COLORS["surface2"], fg=COLORS["text"],
            activebackground=COLORS["accent_hover"], activeforeground="white",
            relief=tk.FLAT, padx=12, pady=6, cursor="hand2",
            font=("Segoe UI", 9), bd=0,
        )

    def _build_ui(self):
        menubar = tk.Menu(self, bg=COLORS["surface"], fg=COLORS["text"],
                          activebackground=COLORS["accent"], activeforeground="white", tearoff=0)
        self.config(menu=menubar)
        file_m = tk.Menu(menubar, tearoff=0, bg=COLORS["surface"], fg=COLORS["text"], activebackground=COLORS["accent"])
        menubar.add_cascade(label="File", menu=file_m)
        file_m.add_command(label="Open...\tCtrl+O", command=self.open_file)
        file_m.add_command(label="Save\tCtrl+S", command=self.save_pdf)
        file_m.add_command(label="Save As...\tCtrl+Shift+S", command=self.save_as_pdf)
        file_m.add_separator()
        file_m.add_command(label="Exit", command=self._on_close)

        edit_m = tk.Menu(menubar, tearoff=0, bg=COLORS["surface"], fg=COLORS["text"], activebackground=COLORS["accent"])
        menubar.add_cascade(label="Edit", menu=edit_m)
        edit_m.add_command(label="Undo\tCtrl+Z", command=self.undo)
        edit_m.add_command(label="Redo\tCtrl+Y", command=self.redo)
        edit_m.add_command(label="Search...\tCtrl+F", command=self.show_search)

        view_m = tk.Menu(menubar, tearoff=0, bg=COLORS["surface"], fg=COLORS["text"], activebackground=COLORS["accent"])
        menubar.add_cascade(label="View", menu=view_m)
        view_m.add_command(label="Zoom In", command=lambda: self.set_zoom(self.zoom * 1.25))
        view_m.add_command(label="Zoom Out", command=lambda: self.set_zoom(self.zoom / 1.25))
        view_m.add_command(label="Fit Page", command=self.fit_page)
        view_m.add_command(label="Fit Width", command=self.fit_width)
        view_m.add_command(label="Actual Size (100%)", command=lambda: self.set_zoom(1.0))

        tools_m = tk.Menu(menubar, tearoff=0, bg=COLORS["surface"], fg=COLORS["text"], activebackground=COLORS["accent"])
        menubar.add_cascade(label="Tools", menu=tools_m)
        tools_m.add_command(label="Draw Signature...", command=self.start_signature)
        tools_m.add_command(label="Copy Sign (select area)...", command=self.start_copy_sign_region)
        tools_m.add_separator()
        tools_m.add_command(label="Color Fill Box (Hide Text)...", command=self.start_fill_box)
        tools_m.add_command(label="Pick Fill Color from Page...", command=self.start_pick_color)
        tools_m.add_command(label="Reset Fill Color (auto background)", command=self.reset_fill_color)
        tools_m.add_separator()
        tools_m.add_command(label="Text Color...", command=self.choose_text_color)
        tools_m.add_command(label="বাংলা কিবোর্ড...", command=self._open_bangla_kb_standalone)
        tools_m.add_separator()
        tools_m.add_command(label="Screenshot", command=self.take_screenshot)
        tools_m.add_command(label="OCR Image Text...", command=self.ocr_current_page)
        tools_m.add_command(label="Set as Default PDF Viewer...", command=self.show_default_viewer_help)

        help_m = tk.Menu(menubar, tearoff=0, bg=COLORS["surface"], fg=COLORS["text"], activebackground=COLORS["accent"])
        menubar.add_cascade(label="Help", menu=help_m)
        help_m.add_command(label="About", command=self.show_about)

        toolbar = tk.Frame(self, bg=COLORS["toolbar"], height=48)
        toolbar.pack(side=tk.TOP, fill=tk.X)
        toolbar.pack_propagate(False)
        inner = tk.Frame(toolbar, bg=COLORS["toolbar"])
        inner.pack(side=tk.LEFT, padx=8, pady=6)
        for text, cmd in [("Open", self.open_file), ("Save", self.save_pdf), ("Save As", self.save_as_pdf)]:
            self._make_tool_btn(inner, text, cmd).pack(side=tk.LEFT, padx=3)
        tk.Frame(inner, width=12, bg=COLORS["toolbar"]).pack(side=tk.LEFT)
        for text, cmd in [("Undo", self.undo), ("Search", self.show_search)]:
            self._make_tool_btn(inner, text, cmd).pack(side=tk.LEFT, padx=3)
        tk.Frame(inner, width=8, bg=COLORS["toolbar"]).pack(side=tk.LEFT)
        for text, cmd in [("Sign", self.start_signature), ("Copy Sign", self.start_copy_sign_region),
                          ("Fill Box", self.start_fill_box), ("Screenshot", self.take_screenshot)]:
            self._make_tool_btn(inner, text, cmd).pack(side=tk.LEFT, padx=3)
        tk.Frame(inner, width=12, bg=COLORS["toolbar"]).pack(side=tk.LEFT)
        self._make_tool_btn(inner, "<", self.prev_page).pack(side=tk.LEFT, padx=2)
        self._make_tool_btn(inner, ">", self.next_page).pack(side=tk.LEFT, padx=2)
        tk.Label(inner, text="Page", bg=COLORS["toolbar"], fg=COLORS["text_dim"], font=("Segoe UI", 9)).pack(side=tk.LEFT, padx=(8, 2))
        self.page_var = tk.StringVar(value="1")
        page_entry = tk.Entry(inner, textvariable=self.page_var, width=4, bg=COLORS["surface2"], fg=COLORS["text"],
                              insertbackground=COLORS["text"], relief=tk.FLAT, font=("Segoe UI", 10))
        page_entry.pack(side=tk.LEFT)
        page_entry.bind("<Return>", self._goto_page)
        self.page_count_lbl = tk.Label(inner, text="/ 0", bg=COLORS["toolbar"], fg=COLORS["text_dim"], font=("Segoe UI", 9))
        self.page_count_lbl.pack(side=tk.LEFT, padx=4)
        tk.Frame(inner, width=12, bg=COLORS["toolbar"]).pack(side=tk.LEFT)
        self._make_tool_btn(inner, "-", lambda: self.set_zoom(self.zoom / 1.25)).pack(side=tk.LEFT)
        self.zoom_lbl = tk.Label(inner, text="100%", bg=COLORS["toolbar"], fg=COLORS["text"], font=("Segoe UI", 9), width=5)
        self.zoom_lbl.pack(side=tk.LEFT, padx=4)
        self._make_tool_btn(inner, "+", lambda: self.set_zoom(self.zoom * 1.25)).pack(side=tk.LEFT)
        self._make_tool_btn(inner, "Fit", self.fit_page).pack(side=tk.LEFT, padx=4)

        self.drop_hint = tk.Label(self, text="  Open PDF / Image / Word file  •  Double-click text to edit in place",
                                  bg=COLORS["surface2"], fg=COLORS["text_dim"], font=("Segoe UI", 10), pady=8)
        self.drop_hint.pack(side=tk.TOP, fill=tk.X)

        main = tk.Frame(self, bg=COLORS["canvas_bg"])
        main.pack(fill=tk.BOTH, expand=True)
        self.canvas = tk.Canvas(main, bg=COLORS["canvas_bg"], highlightthickness=0)
        hsb = ttk.Scrollbar(main, orient=tk.HORIZONTAL, command=self.canvas.xview)
        vsb = ttk.Scrollbar(main, orient=tk.VERTICAL, command=self.canvas.yview)
        self.canvas.configure(xscrollcommand=hsb.set, yscrollcommand=vsb.set)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        hsb.pack(side=tk.BOTTOM, fill=tk.X)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.canvas.bind("<Button-1>", self._on_canvas_click)
        self.canvas.bind("<Double-Button-1>", self._on_double_click)
        self.canvas.bind("<B1-Motion>", self._on_canvas_drag)
        self.canvas.bind("<ButtonRelease-1>", self._on_canvas_release)
        self.canvas.bind("<Configure>", self._on_canvas_configure)
        self.canvas.bind("<MouseWheel>", self._on_mousewheel)
        try:
            import windnd
            windnd.hook_dropfiles(self, func=self._on_windnd_drop)
            windnd.hook_dropfiles(self.canvas, func=self._on_windnd_drop)
        except Exception:
            pass
        self.canvas.bind("<Button-4>", lambda e: self.canvas.yview_scroll(-1, "units"))
        self.canvas.bind("<Button-5>", lambda e: self.canvas.yview_scroll(1, "units"))

        self.status = tk.Label(self, text=f"{APP_NAME} v{APP_VERSION}  •  Made by {APP_STUDIO}",
                               bg=COLORS["surface"], fg=COLORS["text_dim"], font=("Segoe UI", 9),
                               anchor=tk.W, padx=10, pady=4)
        self.status.pack(side=tk.BOTTOM, fill=tk.X)

        self.search_frame = tk.Frame(self, bg=COLORS["surface2"])
        tk.Label(self.search_frame, text="Find:", bg=COLORS["surface2"], fg=COLORS["text"], font=("Segoe UI", 10)).pack(side=tk.LEFT, padx=8)
        self.search_var = tk.StringVar()
        self.search_entry = tk.Entry(self.search_frame, textvariable=self.search_var, width=28,
                                     bg=COLORS["surface"], fg=COLORS["text"], insertbackground=COLORS["text"],
                                     relief=tk.FLAT, font=("Segoe UI", 10))
        self.search_entry.pack(side=tk.LEFT, padx=4, pady=6)
        self.search_entry.bind("<Return>", lambda e: self.do_search())
        self._make_tool_btn(self.search_frame, "Next", self.find_next).pack(side=tk.LEFT, padx=2)
        self._make_tool_btn(self.search_frame, "Prev", self.find_prev).pack(side=tk.LEFT, padx=2)
        self._make_tool_btn(self.search_frame, "Close", self.hide_search).pack(side=tk.LEFT, padx=6)
        self.search_status = tk.Label(self.search_frame, text="", bg=COLORS["surface2"], fg=COLORS["text_dim"], font=("Segoe UI", 9))
        self.search_status.pack(side=tk.LEFT, padx=8)

    def _bind_shortcuts(self):
        self.bind("<Control-o>", lambda e: self.open_file())
        self.bind("<Control-s>", lambda e: self.save_pdf())
        self.bind("<Control-S>", lambda e: self.save_as_pdf())
        self.bind("<Control-Shift-S>", lambda e: self.save_as_pdf())
        self.bind("<Control-Shift-s>", lambda e: self.save_as_pdf())
        self.bind("<Control-f>", lambda e: self.show_search())
        self.bind("<Control-z>", lambda e: self.undo())
        self.bind("<Control-y>", lambda e: self.redo())
        self.bind("<Escape>", lambda e: self._cancel_ops())
        self.bind("<Prior>", lambda e: self.prev_page())
        self.bind("<Next>", lambda e: self.next_page())

    def open_file(self):
        if self.pdf.dirty:
            if not messagebox.askyesno("Unsaved", "Discard unsaved changes?"):
                return
        path = filedialog.askopenfilename(
            title="Open File",
            filetypes=[
                ("All supported", "*.pdf;*.png;*.jpg;*.jpeg;*.bmp;*.gif;*.webp;*.docx"),
                ("PDF", "*.pdf"),
                ("Images", "*.png;*.jpg;*.jpeg;*.bmp;*.gif;*.webp;*.tiff"),
                ("Word", "*.docx"),
                ("All files", "*.*"),
            ],
        )
        if path:
            self._load_path(path)

    def _load_path(self, path):
        try:
            ok = self.pdf.open(path)
            if not ok:
                pwd = simpledialog.askstring("Password", "File is encrypted. Enter password:", show="*")
                if pwd is None:
                    return
                ok = self.pdf.open(path, password=pwd)
                if not ok:
                    messagebox.showerror("Error", "Incorrect password or cannot open.")
                    return
            self.current_page = 0
            self.zoom = DEFAULT_ZOOM
            self.highlight_rect = None
            self.selected_span = None
            self._cancel_inline_edit()
            self._update_page_ui()
            self.render_page()
            self.fit_width()
            name = os.path.basename(path)
            kind = self.pdf.source_type.upper()
            self.status.config(text=f"Opened ({kind}): {name}")
            try:
                self.drop_hint.pack_forget()
            except Exception:
                pass
        except Exception as e:
            messagebox.showerror("Open failed", str(e))

    def save_pdf(self):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.", parent=self)
            return
        if not self.pdf.path or not str(self.pdf.path).lower().endswith(".pdf"):
            return self.save_as_pdf()
        try:
            self.pdf.save()
            self.status.config(text=f"Saved: {os.path.basename(self.pdf.path)}")
            messagebox.showinfo("Saved", f"Saved to:\n{self.pdf.path}", parent=self)
        except Exception as e:
            messagebox.showerror("Save failed", str(e), parent=self)

    def save_as_pdf(self):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.", parent=self)
            return
        # Ensure window is focused so native dialog appears on Windows
        try:
            self.lift()
            self.focus_force()
            self.update_idletasks()
        except Exception:
            pass
        initial_name = Path(self.pdf.path or "document").stem + ".pdf"
        initial_dir = None
        try:
            if self.pdf.path and Path(self.pdf.path).parent.exists():
                initial_dir = str(Path(self.pdf.path).parent)
            else:
                initial_dir = str(Path.home() / "Documents")
        except Exception:
            initial_dir = str(Path.home())
        path = filedialog.asksaveasfilename(
            parent=self,
            title="Save As — choose name and folder",
            defaultextension=".pdf",
            filetypes=[
                ("PDF files", "*.pdf"),
                ("All files", "*.*"),
            ],
            initialfile=initial_name,
            initialdir=initial_dir,
            confirmoverwrite=True,
        )
        if not path:
            self.status.config(text="Save As cancelled")
            return
        if not path.lower().endswith(".pdf"):
            path = path + ".pdf"
        try:
            self.pdf.save(path)
            self.status.config(text=f"Saved as: {os.path.basename(path)}")
            messagebox.showinfo("Saved", f"File saved as:\n{path}", parent=self)
        except Exception as e:
            messagebox.showerror("Save As failed", str(e), parent=self)

    def undo(self):
        self._cancel_inline_edit()
        if self.pdf.undo():
            self.highlight_rect = None
            self.selected_span = None
            self.render_page()
            self.status.config(text="Undo")
        else:
            self.status.config(text="Nothing to undo")

    def redo(self):
        self._cancel_inline_edit()
        if self.pdf.redo():
            self.highlight_rect = None
            self.selected_span = None
            self.render_page()
            self.status.config(text="Redo")
        else:
            self.status.config(text="Nothing to redo")

    def prev_page(self):
        if self.current_page > 0:
            self._cancel_inline_edit()
            self.current_page -= 1
            self.highlight_rect = None
            self.selected_span = None
            self._update_page_ui()
            self.render_page()

    def next_page(self):
        if self.pdf.doc and self.current_page < self.pdf.page_count() - 1:
            self._cancel_inline_edit()
            self.current_page += 1
            self.highlight_rect = None
            self.selected_span = None
            self._update_page_ui()
            self.render_page()

    def _goto_page(self, event=None):
        try:
            p = int(self.page_var.get()) - 1
            if 0 <= p < self.pdf.page_count():
                self._cancel_inline_edit()
                self.current_page = p
                self.highlight_rect = None
                self.selected_span = None
                self._update_page_ui()
                self.render_page()
        except ValueError:
            pass

    def set_zoom(self, z):
        self._cancel_inline_edit()
        self.zoom = max(0.25, min(4.0, z))
        self.zoom_lbl.config(text=f"{int(self.zoom * 100)}%")
        self.render_page()

    def fit_page(self):
        if not self.pdf.doc:
            return
        page = self.pdf.get_page(self.current_page)
        if not page:
            return
        cw = max(self.canvas.winfo_width(), 100)
        ch = max(self.canvas.winfo_height(), 100)
        rect = page.rect
        self.set_zoom(min((cw - 24) / rect.width, (ch - 24) / rect.height))

    def fit_width(self):
        if not self.pdf.doc:
            return
        page = self.pdf.get_page(self.current_page)
        if not page:
            return
        cw = max(self.canvas.winfo_width(), 100)
        self.set_zoom((cw - 24) / page.rect.width)

    def _update_page_ui(self):
        n = self.pdf.page_count()
        self.page_count_lbl.config(text=f"/ {n}")
        self.page_var.set(str(self.current_page + 1))

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def render_page(self):
        self.canvas.delete("all")
        self.photo = None
        if not self.pdf.doc:
            self._show_dashboard()
            return
        self._stop_dashboard()
        page = self.pdf.get_page(self.current_page)
        if not page:
            return
        try:
            mat = fitz.Matrix(self.zoom, self.zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            self.photo = ImageTk.PhotoImage(img)
            self.canvas.create_image(0, 0, anchor=tk.NW, image=self.photo, tags="page")
            self.canvas.config(scrollregion=(0, 0, pix.width, pix.height))
            if self.highlight_rect:
                r = self.highlight_rect * self.zoom
                self.canvas.create_rectangle(r.x0, r.y0, r.x1, r.y1, outline="#fbbf24", width=2, tags="highlight")
            if self.selected_span and self.selected_page == self.current_page:
                r = self.selected_span["bbox"] * self.zoom
                self.canvas.create_rectangle(r.x0, r.y0, r.x1, r.y1, outline="#4ade80", width=2, tags="selected")
            if self.placing_signature and self.sig_rect:
                r = self.sig_rect * self.zoom
                self.canvas.create_rectangle(r.x0, r.y0, r.x1, r.y1, outline="#7c5cff", width=2, dash=(4, 2), tags="sigpreview")
            if self.fill_mode and self.fill_rect:
                r = self.fill_rect * self.zoom
                if self.fill_color_manual and self.fill_color:
                    c = self.fill_color
                    fill_hex = "#%02x%02x%02x" % (int(c[0]*255), int(c[1]*255), int(c[2]*255))
                else:
                    fill_hex = "#cccccc"
                self.canvas.create_rectangle(r.x0, r.y0, r.x1, r.y1, outline="#f87171", width=2,
                                            fill=fill_hex, stipple="gray50", tags="fillpreview")
            if self.copy_sign_mode and self.copy_sign_rect:
                r = self.copy_sign_rect * self.zoom
                self.canvas.create_rectangle(r.x0, r.y0, r.x1, r.y1, outline="#22d3ee", width=2,
                                            dash=(4, 2), tags="copysignpreview")
            if self.screenshot_mode and self.screenshot_rect:
                r = self.screenshot_rect * self.zoom
                self.canvas.create_rectangle(r.x0, r.y0, r.x1, r.y1, outline="#fbbf24", width=2,
                                            dash=(5, 3), tags="sspreview")
        except Exception as e:
            self.status.config(text=f"Render error: {e}")

    def _canvas_to_pdf(self, x, y):
        return fitz.Point(x / self.zoom, y / self.zoom)

    def _on_canvas_click(self, event):
        if not self.pdf.doc:
            return
        if self._edit_entry:
            self._finish_inline_edit()
            return
        cx = self.canvas.canvasx(event.x)
        cy = self.canvas.canvasy(event.y)
        pt = self._canvas_to_pdf(cx, cy)

        # Color picker mode
        if self.pick_color_mode:
            self._sample_color_at(cx, cy)
            return

        # Screenshot select area
        if self.screenshot_mode:
            self._drag_start = pt
            self.screenshot_rect = fitz.Rect(pt, pt)
            return

        # Copy sign region mode
        if self.copy_sign_mode:
            self._drag_start = pt
            self.copy_sign_rect = fitz.Rect(pt, pt)
            return

        # Fill box mode – start drag
        if self.fill_mode:
            self._drag_start = pt
            self.fill_rect = fitz.Rect(pt, pt)
            return

        if self.placing_signature and self.signature_img:
            self._drag_start = pt
            self.sig_rect = fitz.Rect(pt, pt)
            return
        spans = self.pdf.get_text_spans(self.current_page)
        hit = None
        for s in spans:
            if s["bbox"].contains(pt):
                hit = s
                break
        if hit:
            self.selected_span = hit
            self.selected_page = self.current_page
            self.status.config(text=f'Selected: "{hit["text"][:50]}"  —  Double-click to edit in place')
            self.render_page()
        else:
            self.selected_span = None
            self.render_page()

    def _on_double_click(self, event):
        if self.selected_span and self.selected_page == self.current_page:
            self._start_inline_edit()

    def _start_inline_edit(self):
        if not self.selected_span:
            return
        self._cancel_inline_edit()
        span = self.selected_span
        r = span["bbox"] * self.zoom
        # Format bar above the text
        self._fmt_frame = tk.Frame(self.canvas, bg="#2a2a3c", padx=4, pady=2)
        tk.Label(self._fmt_frame, text="Size:", bg="#2a2a3c", fg="#e8e8f0", font=("Segoe UI", 8)).pack(side=tk.LEFT)
        self._fmt_size = tk.StringVar(value=str(int(span["size"])))
        size_entry = tk.Entry(self._fmt_frame, textvariable=self._fmt_size, width=4,
                              bg="#34344a", fg="#e8e8f0", relief=tk.FLAT, font=("Segoe UI", 9))
        size_entry.pack(side=tk.LEFT, padx=2)
        self._fmt_bold = tk.BooleanVar(value=bool(span.get("flags", 0) & 16))
        self._fmt_italic = tk.BooleanVar(value=bool(span.get("flags", 0) & 2))
        tk.Checkbutton(self._fmt_frame, text="B", variable=self._fmt_bold, bg="#2a2a3c", fg="#e8e8f0",
                       selectcolor="#7c5cff", font=("Segoe UI", 9, "bold"),
                       activebackground="#2a2a3c").pack(side=tk.LEFT, padx=2)
        tk.Checkbutton(self._fmt_frame, text="I", variable=self._fmt_italic, bg="#2a2a3c", fg="#e8e8f0",
                       selectcolor="#7c5cff", font=("Segoe UI", 9, "italic"),
                       activebackground="#2a2a3c").pack(side=tk.LEFT, padx=2)
        tk.Button(self._fmt_frame, text="বাংলা", command=self._open_bangla_kb,
                  bg="#0ea5e9", fg="white", relief=tk.FLAT, font=("Segoe UI", 8), padx=6).pack(side=tk.LEFT, padx=4)
        tk.Button(self._fmt_frame, text="Apply", command=self._finish_inline_edit,
                  bg="#7c5cff", fg="white", relief=tk.FLAT, font=("Segoe UI", 8), padx=6).pack(side=tk.LEFT, padx=4)
        tk.Button(self._fmt_frame, text="X", command=self._cancel_inline_edit,
                  bg="#34344a", fg="#e8e8f0", relief=tk.FLAT, font=("Segoe UI", 8), padx=4).pack(side=tk.LEFT)
        self.canvas.create_window(r.x0, max(0, r.y0 - 28), window=self._fmt_frame, anchor=tk.NW, tags="editentry")
        self._edit_entry = tk.Entry(
            self.canvas,
            font=("Segoe UI", max(9, int(span["size"] * self.zoom * 0.75))),
            bg="#fffde7", fg="#111", relief=tk.SOLID, bd=1, insertbackground="#111",
        )
        self._edit_entry.insert(0, span["text"])
        self._edit_entry.select_range(0, tk.END)
        self._edit_span = span
        self.canvas.create_window(r.x0, r.y0, window=self._edit_entry, anchor=tk.NW,
                                  width=max(r.width, 120), height=max(r.height + 4, 22), tags="editentry")
        self._edit_entry.focus_set()
        self._edit_entry.bind("<Return>", lambda e: self._finish_inline_edit())
        self._edit_entry.bind("<Escape>", lambda e: self._cancel_inline_edit())
        self.status.config(text="Edit text + set Size / Bold / Italic — Apply or Enter")

    def _open_bangla_kb(self):
        """Open on-screen Bangla keyboard targeting the inline edit Entry."""
        if not self._edit_entry:
            messagebox.showinfo("Edit first", "Double-click text to edit, then open বাংলা keyboard.")
            return

        def insert_char(ch):
            if not self._edit_entry:
                return
            if ch == "BACK":
                try:
                    self._edit_entry.delete(len(self._edit_entry.get()) - 1, tk.END)
                except Exception:
                    pass
                return
            self._edit_entry.insert(tk.INSERT, ch)
            self._edit_entry.focus_set()

        BanglaKeyboard(self, on_char=insert_char, on_enter=self._finish_inline_edit)

    def _finish_inline_edit(self):
        if not self._edit_entry or not self._edit_span:
            self._cancel_inline_edit()
            return
        new_text = self._edit_entry.get()
        span = self._edit_span
        page = self.selected_page
        try:
            fs = float(self._fmt_size.get())
        except Exception:
            fs = float(span["size"])
        bold = bool(self._fmt_bold.get()) if hasattr(self, "_fmt_bold") else False
        italic = bool(self._fmt_italic.get()) if hasattr(self, "_fmt_italic") else False
        self._cancel_inline_edit()
        if not new_text.strip():
            return
        ok, msg = self.pdf.replace_span(page, span, new_text, fontsize=fs, bold=bold, italic=italic)
        self.selected_span = None
        if ok:
            self.render_page()
            self.status.config(text=msg)
        else:
            messagebox.showerror("Edit failed", msg)

    def _cancel_inline_edit(self):
        if self._edit_entry:
            try:
                self._edit_entry.destroy()
            except Exception:
                pass
        if hasattr(self, "_fmt_frame") and self._fmt_frame:
            try:
                self._fmt_frame.destroy()
            except Exception:
                pass
            self._fmt_frame = None
        self._edit_entry = None
        self._edit_span = None
        try:
            self.canvas.delete("editentry")
        except Exception:
            pass

    def _on_canvas_drag(self, event):
        if not self._drag_start:
            return
        cx = self.canvas.canvasx(event.x)
        cy = self.canvas.canvasy(event.y)
        pt = self._canvas_to_pdf(cx, cy)
        if self.screenshot_mode:
            self.screenshot_rect = fitz.Rect(self._drag_start, pt)
            self.screenshot_rect.normalize()
            self.render_page()
            return
        if self.copy_sign_mode:
            self.copy_sign_rect = fitz.Rect(self._drag_start, pt)
            self.copy_sign_rect.normalize()
            self.render_page()
            return
        if self.fill_mode:
            self.fill_rect = fitz.Rect(self._drag_start, pt)
            self.fill_rect.normalize()
            self.render_page()
            return
        if self.placing_signature:
            self.sig_rect = fitz.Rect(self._drag_start, pt)
            self.sig_rect.normalize()
            self.render_page()

    def _on_canvas_release(self, event):
        if self.screenshot_mode and self.screenshot_rect:
            self._screenshot_selected_area()
            return
        if self.copy_sign_mode and self.copy_sign_rect:
            self._capture_sign_region()
            return
        if self.fill_mode and self.fill_rect:
            if self.fill_rect.width < 5 or self.fill_rect.height < 5:
                self.status.config(text="Box too small")
                return
            self._apply_fill_box()
            return
        if self.placing_signature and self.sig_rect and self.signature_img:
            if self.sig_rect.width < 10 or self.sig_rect.height < 5:
                self.status.config(text="Signature area too small")
                return
            ok = self.pdf.insert_signature_image(self.current_page, self.signature_img, self.sig_rect)
            if ok:
                self.status.config(text="Signature placed")
                self.placing_signature = False
                self.signature_img = None
                self.sig_rect = None
                self._drag_start = None
                self.render_page()
            else:
                messagebox.showerror("Error", "Could not insert signature")

    def show_search(self):
        self.search_frame.pack(side=tk.TOP, fill=tk.X)
        self.search_entry.focus_set()
        self.search_entry.select_range(0, tk.END)

    def hide_search(self):
        self.search_frame.pack_forget()
        self.highlight_rect = None
        self.render_page()

    def do_search(self):
        q = self.search_var.get().strip()
        if not q or not self.pdf.doc:
            return
        self.search_results = self.pdf.search(q)
        self.search_index = -1
        if not self.search_results:
            self.search_status.config(text="No matches")
            self.highlight_rect = None
            self.render_page()
        else:
            self.search_status.config(text=f"{len(self.search_results)} found")
            self.find_next()

    def find_next(self):
        if not self.search_results:
            self.do_search()
            return
        self.search_index = (self.search_index + 1) % len(self.search_results)
        self._goto_search_hit()

    def find_prev(self):
        if not self.search_results:
            return
        self.search_index = (self.search_index - 1) % len(self.search_results)
        self._goto_search_hit()

    def _goto_search_hit(self):
        page_idx, rect = self.search_results[self.search_index]
        self.current_page = page_idx
        self.highlight_rect = rect
        self._update_page_ui()
        self.render_page()
        self.search_status.config(text=f"{self.search_index + 1} / {len(self.search_results)}")

    def start_signature(self):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.")
            return
        def on_done(img):
            if img is None:
                return
            self.signature_img = img
            self.placing_signature = True
            self.sig_rect = None
            self._drag_start = None
            self.status.config(text="Drag a rectangle on the page to place the signature")
        SignatureDialog(self, on_done)

    def take_screenshot(self):
        if not self.pdf.doc:
            messagebox.showinfo("No page", "Open a file first.")
            return
        # Choice dialog
        win = tk.Toplevel(self)
        win.title("Screenshot")
        win.configure(bg=COLORS["surface"])
        win.transient(self)
        win.grab_set()
        win.resizable(False, False)
        tk.Label(win, text="Choose screenshot type", bg=COLORS["surface"], fg=COLORS["text"],
                 font=("Segoe UI", 12, "bold")).pack(padx=20, pady=(16, 8))
        bf = tk.Frame(win, bg=COLORS["surface"])
        bf.pack(padx=20, pady=(4, 16))

        def do_page():
            win.destroy()
            self._screenshot_full_page()

        def do_area():
            win.destroy()
            self.screenshot_mode = True
            self.fill_mode = False
            self.copy_sign_mode = False
            self.placing_signature = False
            self.pick_color_mode = False
            self.screenshot_rect = None
            self._drag_start = None
            try:
                self.canvas.config(cursor="crosshair")
            except Exception:
                pass
            self.status.config(text="Screenshot area: drag a rectangle on the page")

        tk.Button(bf, text="📄  Current Page", command=do_page,
                  bg=COLORS["accent"], fg="white", relief=tk.FLAT, padx=18, pady=10,
                  font=("Segoe UI", 10), cursor="hand2").pack(side=tk.LEFT, padx=6)
        tk.Button(bf, text="⬚  Select Area", command=do_area,
                  bg=COLORS["surface2"], fg=COLORS["text"], relief=tk.FLAT, padx=18, pady=10,
                  font=("Segoe UI", 10), cursor="hand2").pack(side=tk.LEFT, padx=6)
        tk.Button(win, text="Cancel", command=win.destroy,
                  bg=COLORS["surface2"], fg=COLORS["text_dim"], relief=tk.FLAT, padx=12, pady=6,
                  font=("Segoe UI", 9), cursor="hand2").pack(pady=(0, 14))
        win.update_idletasks()
        x = self.winfo_rootx() + (self.winfo_width() - win.winfo_width()) // 2
        y = self.winfo_rooty() + (self.winfo_height() - win.winfo_height()) // 2
        win.geometry(f"+{x}+{y}")

    def _screenshot_full_page(self):
        try:
            page = self.pdf.get_page(self.current_page)
            if not page:
                return
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            ScreenshotPopup(self, img)
        except Exception as e:
            messagebox.showerror("Screenshot failed", str(e))

    def _screenshot_selected_area(self):
        rect = self.screenshot_rect
        self.screenshot_mode = False
        self.screenshot_rect = None
        self._drag_start = None
        try:
            self.canvas.config(cursor="")
        except Exception:
            pass
        if not rect or rect.width < 5 or rect.height < 5:
            self.status.config(text="Screenshot area too small")
            self.render_page()
            return
        try:
            page = self.pdf.get_page(self.current_page)
            if not page:
                return
            mat = fitz.Matrix(2.5, 2.5)
            pix = page.get_pixmap(matrix=mat, clip=rect, alpha=False)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            ScreenshotPopup(self, img)
            self.status.config(text="Area screenshot ready")
            self.render_page()
        except Exception as e:
            messagebox.showerror("Screenshot failed", str(e))

    def _cancel_ops(self):
        self._cancel_inline_edit()
        self.placing_signature = False
        self.signature_img = None
        self.sig_rect = None
        self.fill_mode = False
        self.pick_color_mode = False
        self.copy_sign_mode = False
        self.screenshot_mode = False
        self.fill_rect = None
        self.copy_sign_rect = None
        self.screenshot_rect = None
        self._drag_start = None
        try:
            self.canvas.config(cursor="")
        except Exception:
            pass
        self.hide_search()
        self.status.config(text="Cancelled")


    def start_pick_color(self):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.")
            return
        self.pick_color_mode = True
        self.fill_mode = False
        self.placing_signature = False
        self.status.config(text="Click anywhere on the page to pick that color")

    def _sample_color_at(self, canvas_x, canvas_y):
        """Sample RGB from current rendered page pixmap."""
        try:
            page = self.pdf.get_page(self.current_page)
            if not page:
                return
            mat = fitz.Matrix(self.zoom, self.zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            x = int(canvas_x)
            y = int(canvas_y)
            if 0 <= x < pix.width and 0 <= y < pix.height:
                # pix.samples is RGB
                i = (y * pix.width + x) * 3
                r = pix.samples[i] / 255.0
                g = pix.samples[i + 1] / 255.0
                b = pix.samples[i + 2] / 255.0
                self.fill_color = (r, g, b)
                self.fill_color_manual = True
                self.pick_color_mode = False
                hexc = "#%02x%02x%02x" % (int(r*255), int(g*255), int(b*255))
                self.status.config(text=f"Fill color picked {hexc} — use Fill Box")
                messagebox.showinfo("Color Picked", f"Fill color: {hexc}\n\nNow use Fill Box and drag a rectangle.")
            else:
                self.status.config(text="Click inside the page")
        except Exception as e:
            self.status.config(text=f"Pick failed: {e}")
            self.pick_color_mode = False

    def start_fill_box(self):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.")
            return
        self.fill_mode = True
        self.pick_color_mode = False
        self.placing_signature = False
        self.copy_sign_mode = False
        self.fill_rect = None
        self._drag_start = None
        if self.fill_color_manual and self.fill_color:
            c = self.fill_color
            hexc = "#%02x%02x%02x" % (int(c[0]*255), int(c[1]*255), int(c[2]*255))
            self.status.config(text=f"Fill Box (manual {hexc}): drag rectangle")
        else:
            self.status.config(text="Fill Box (auto match PDF background): drag rectangle")

    def reset_fill_color(self):
        self.fill_color = None
        self.fill_color_manual = False
        self.status.config(text="Fill color reset — will auto-match PDF background")

    def choose_text_color(self):
        initial = (
            int(self.text_color[0] * 255),
            int(self.text_color[1] * 255),
            int(self.text_color[2] * 255),
        )
        result = colorchooser.askcolor(color=initial, title="Choose text color")
        if result and result[0]:
            r, g, b = result[0]
            self.text_color = (r / 255.0, g / 255.0, b / 255.0)
            self.status.config(text=f"Text color set to #{int(r):02x}{int(g):02x}{int(b):02x}")

    def _apply_fill_box(self):
        rect = self.fill_rect
        if not rect:
            return
        # Keep a copy — doc changes must not lose coordinates
        saved_rect = fitz.Rect(rect)
        # Auto-sample background unless user manually picked a fill color
        if self.fill_color_manual and self.fill_color is not None:
            use_color = self.fill_color
        else:
            use_color = self.pdf.sample_rect_background(self.current_page, saved_rect)
        ok = self.pdf.fill_rectangle(self.current_page, saved_rect, use_color)
        self.fill_mode = False
        self.fill_rect = None
        self._drag_start = None
        if not ok:
            messagebox.showerror("Error", "Could not apply fill box")
            self.render_page()
            return
        self.render_page()
        if messagebox.askyesno("Write text?", "Area covered.\n\nDo you want to type new text on this area?"):
            text = self._ask_text_with_bangla("Enter text (English or use বাংলা keyboard):")
            if text and text.strip():
                ok2 = self.pdf.insert_text_in_rect(
                    self.current_page, saved_rect, text.strip(),
                    fontsize=None, color=self.text_color,
                )
                self.render_page()
                if ok2:
                    self.status.config(text="Fill box + new text applied")
                else:
                    messagebox.showerror("Text failed", "Could not insert text. Try a larger box.")
                    self.status.config(text="Fill applied but text insert failed")
            else:
                self.status.config(text="Fill box applied (text hidden)")
        else:
            self.status.config(text="Fill box applied (text hidden)")

    def copy_signature_from_pdf(self):
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.")
            return
        images = self.pdf.extract_page_images(self.current_page)
        if not images:
            messagebox.showinfo("No images", "No embedded images found on this page.\nSignatures drawn as vector may not appear.")
            return
        # Simple chooser dialog
        win = tk.Toplevel(self)
        win.title("Copy Signature from Page")
        win.configure(bg=COLORS["surface"])
        win.transient(self)
        win.grab_set()
        tk.Label(win, text="Select an image to use as signature:", bg=COLORS["surface"],
                 fg=COLORS["text"], font=("Segoe UI", 10)).pack(padx=12, pady=8)
        frame = tk.Frame(win, bg=COLORS["surface"])
        frame.pack(padx=12, pady=4)
        self._sig_choice_imgs = []
        for idx, (xref, img) in enumerate(images[:12]):
            thumb = img.copy()
            thumb.thumbnail((120, 60))
            photo = ImageTk.PhotoImage(thumb)
            self._sig_choice_imgs.append(photo)
            def make_cmd(im=img):
                def cmd():
                    self.signature_img = im.convert("RGBA") if im.mode != "RGBA" else im
                    self.placing_signature = True
                    self.sig_rect = None
                    self._drag_start = None
                    self.status.config(text="Signature copied — drag a rectangle to place it")
                    win.destroy()
                return cmd
            btn = tk.Button(frame, image=photo, command=make_cmd(), bg=COLORS["surface2"], relief=tk.FLAT)
            btn.grid(row=idx // 4, column=idx % 4, padx=4, pady=4)
        tk.Button(win, text="Cancel", command=win.destroy, bg=COLORS["surface2"], fg=COLORS["text"],
                  relief=tk.FLAT, padx=12, pady=6).pack(pady=10)
        win.geometry("560x280")

    def show_default_viewer_help(self):
        messagebox.showinfo(
            "Set as Default PDF Viewer",
            "To open PDFs with One PDF Editor by default on Windows:\n\n"
            "1. Right-click any PDF file\n"
            "2. Choose Open with → Choose another app\n"
            "3. Browse and select OnePDFEditor.exe\n"
            "4. Check Always use this app to open .pdf files\n"
            "5. Click OK\n\n"
            "Or: Settings → Apps → Default apps → PDF → One PDF Editor\n\n"
            "This must be done once on your PC after you have the .exe.",
        )


    def ocr_current_page(self):
        """Try to OCR current page so image text becomes selectable/editable."""
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.")
            return
        page = self.pdf.get_page(self.current_page)
        if not page:
            return
        try:
            self.status.config(text="Running OCR… please wait")
            self.update_idletasks()
            self.pdf.save_state()
            # PyMuPDF OCR (requires Tesseract installed on the PC)
            tp = page.get_textpage_ocr(dpi=200, full=True)
            text = page.get_text("text", textpage=tp)
            if not text.strip():
                messagebox.showinfo(
                    "OCR result",
                    "No text detected.\n\n"
                    "For OCR you need Tesseract installed on Windows:\n"
                    "https://github.com/UB-Mannheim/tesseract/wiki\n\n"
                    "Without Tesseract, you can still:\n"
                    "• Cover areas with Fill Box\n"
                    "• Type new text on top\n"
                    "• Add signatures",
                )
                self.status.config(text="OCR: no text found (is Tesseract installed?)")
                return
            # Re-insert detected text as real PDF text (simple full-page overlay approach)
            # Better: user can now search; for editing, spans from OCR textpage help
            messagebox.showinfo(
                "OCR done",
                f"Detected text ({len(text)} chars).\n\n"
                "You can now Search the text.\n"
                "For precise edit: use Fill Box to cover old glyphs, then type new text.\n\n"
                "Preview of detected text:\n" + text[:300],
            )
            self.status.config(text="OCR finished — search works on detected text")
            self.render_page()
        except Exception as e:
            messagebox.showinfo(
                "OCR not available",
                "Could not run OCR.\n\n"
                "Install Tesseract for Windows, then restart the app.\n"
                "Download: https://github.com/UB-Mannheim/tesseract/wiki\n\n"
                f"Error: {e}\n\n"
                "Without OCR you can still hide text (Fill Box) and write new text.",
            )
            self.status.config(text="OCR unavailable — install Tesseract for image text edit")


    def start_copy_sign_region(self):
        """Plus-cursor: drag a rectangle; capture that area as signature."""
        if not self.pdf.doc:
            messagebox.showinfo("No file", "Open a file first.")
            return
        self.copy_sign_mode = True
        self.fill_mode = False
        self.pick_color_mode = False
        self.placing_signature = False
        self.copy_sign_rect = None
        self._drag_start = None
        try:
            self.canvas.config(cursor="crosshair")
        except Exception:
            pass
        self.status.config(text="Copy Sign: drag a rectangle over the signature area")

    def _capture_sign_region(self):
        rect = self.copy_sign_rect
        if not rect or rect.width < 5 or rect.height < 5:
            self.status.config(text="Area too small")
            return
        try:
            page = self.pdf.get_page(self.current_page)
            # high-res capture of region
            mat = fitz.Matrix(2.5, 2.5)
            pix = page.get_pixmap(matrix=mat, clip=rect, alpha=True)
            mode = "RGBA" if pix.alpha else "RGB"
            img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
            if img.mode != "RGBA":
                img = img.convert("RGBA")
            self.signature_img = img
            self.copy_sign_mode = False
            self.copy_sign_rect = None
            self._drag_start = None
            self.placing_signature = True
            try:
                self.canvas.config(cursor="")
            except Exception:
                pass
            self.status.config(text="Signature captured — now drag where you want to place it")
            self.render_page()
        except Exception as e:
            messagebox.showerror("Copy failed", str(e))
            self.copy_sign_mode = False
            try:
                self.canvas.config(cursor="")
            except Exception:
                pass


    def _show_dashboard(self):
        """Home screen: near-realtime video frame playback + transparent text."""
        try:
            self.page_count_lbl.config(text="/ -")
            self.page_var.set("-")
        except Exception:
            pass
        try:
            self.status.config(text="")
        except Exception:
            pass

        # Load all video frames once (12 fps sequence)
        if not getattr(self, "_dash_photos", None):
            self._dash_photos = []
            roots = [resource_path("assets", "dashboard")]
            if getattr(sys, "frozen", False):
                roots.extend([
                    Path(sys.executable).parent / "assets" / "dashboard",
                    Path(sys.executable).parent / "_internal" / "assets" / "dashboard",
                ])
            frame_files = []
            for root in roots:
                if not root.exists():
                    continue
                # f0001.jpg style from ffmpeg
                found = sorted(root.glob("f*.jpg"))
                if not found:
                    found = sorted(root.glob("bg*.jpg"), key=lambda p: p.name)
                if found:
                    frame_files = found
                    break
            for fp in frame_files:
                try:
                    self._dash_photos.append(Image.open(fp).convert("RGB"))
                except Exception:
                    pass
            self._dash_idx = 0
            self._dash_size = (0, 0)
            self._dash_scaled = []  # cache scaled frames for current canvas size

        self.update_idletasks()
        cw = max(int(self.canvas.winfo_width()), 400)
        ch = max(int(self.canvas.winfo_height()), 300)
        if cw < 50:
            cw, ch = 900, 600

        self._draw_dashboard_frame(cw, ch)
        self._schedule_dashboard_cycle()

    def _ensure_scaled_frames(self, cw, ch):
        """Cover-scale all frames to canvas size (cached until resize)."""
        if self._dash_scaled and getattr(self, "_dash_size", None) == (cw, ch):
            return
        self._dash_size = (cw, ch)
        self._dash_scaled = []
        if not self._dash_photos:
            return
        for im0 in self._dash_photos:
            im = im0.copy()
            iw, ih = im.size
            scale = max(cw / max(iw, 1), ch / max(ih, 1))
            nw, nh = max(1, int(iw * scale)), max(1, int(ih * scale))
            im = im.resize((nw, nh), Image.Resampling.BILINEAR)
            left = max(0, (nw - cw) // 2)
            top = max(0, (nh - ch) // 2)
            im = im.crop((left, top, left + cw, top + ch))
            self._dash_scaled.append(im)

    def _draw_dashboard_frame(self, cw=None, ch=None):
        if cw is None or ch is None:
            cw = max(int(self.canvas.winfo_width()), 400)
            ch = max(int(self.canvas.winfo_height()), 300)
        self.canvas.delete("dash")
        if self._dash_photos:
            self._ensure_scaled_frames(cw, ch)
            if self._dash_scaled:
                im = self._dash_scaled[self._dash_idx % len(self._dash_scaled)]
                self._dash_bg_photo = ImageTk.PhotoImage(im)
                self.canvas.create_image(0, 0, anchor=tk.NW, image=self._dash_bg_photo, tags="dash")
                self.canvas.config(scrollregion=(0, 0, cw, ch), bg="#0b1220")
        else:
            self.canvas.config(bg="#1a2744")
            self.canvas.create_rectangle(0, 0, cw, ch, fill="#1a2744", outline="", tags="dash")

        # Transparent text only
        cx, cy = cw // 2, ch // 2
        line1, line2 = "Drag & Drop", "your file here"
        for dx, dy in ((1, 1), (0, 1)):
            self.canvas.create_text(cx + dx, cy - 18 + dy, text=line1, fill="#000000",
                                    font=("Segoe UI", 28, "bold"), tags="dash")
            self.canvas.create_text(cx + dx, cy + 22 + dy, text=line2, fill="#000000",
                                    font=("Segoe UI", 16), tags="dash")
        self.canvas.create_text(cx, cy - 18, text=line1, fill="#1a1a1a",
                                font=("Segoe UI", 28, "bold"), tags="dash")
        self.canvas.create_text(cx, cy + 22, text=line2, fill="#1a1a1a",
                                font=("Segoe UI", 16), tags="dash")

    def _schedule_dashboard_cycle(self):
        self._stop_dashboard(cancel_only=True)
        # ~12 fps to match source video
        if not self.pdf.doc and self._dash_photos:
            self._dash_job = self.after(83, self._cycle_dashboard)

    def _cycle_dashboard(self):
        self._dash_job = None
        if self.pdf.doc or not self._dash_photos:
            return
        self._dash_idx = (self._dash_idx + 1) % len(self._dash_photos)
        # Only redraw frame (fast path)
        try:
            self._draw_dashboard_frame()
        except Exception:
            pass
        self._schedule_dashboard_cycle()

    def _stop_dashboard(self, cancel_only=False):
        if self._dash_job is not None:
            try:
                self.after_cancel(self._dash_job)
            except Exception:
                pass
            self._dash_job = None

    def _on_drop_files(self, event):
        """Handle drag-and-drop of files onto the window/canvas."""
        try:
            data = event.data
        except Exception:
            data = str(event)
        # tkdnd or windows: {path} or path
        paths = []
        if isinstance(data, str):
            raw = data.strip()
            if raw.startswith("{"):
                # multiple {a} {b}
                import re as _re
                paths = _re.findall(r"\{([^}]+)\}", raw)
                if not paths:
                    paths = [raw.strip("{}")]
            else:
                paths = [raw]
        for p in paths:
            p = p.strip().strip("{}")
            if os.path.isfile(p):
                self._load_path(p)
                break

    def _on_windnd_drop(self, files):
        try:
            for f in files:
                path = f.decode("utf-8") if isinstance(f, bytes) else str(f)
                if os.path.isfile(path):
                    self._load_path(path)
                    break
        except Exception as e:
            self.status.config(text=f"Drop failed: {e}")

    def _on_canvas_configure(self, event):
        if not self.pdf.doc and event.width > 50 and event.height > 50:
            # invalidate scale cache on resize
            self._dash_size = (0, 0)
            self._dash_scaled = []
            if self._dash_job is None:
                self._dash_job = self.after(100, lambda: self._draw_dashboard_frame(event.width, event.height) or self._schedule_dashboard_cycle())

    def _ask_text_with_bangla(self, prompt="Enter text:"):
        """Text entry dialog with optional on-screen Bangla keyboard."""
        win = tk.Toplevel(self)
        win.title("Type text")
        win.configure(bg=COLORS["surface"])
        win.transient(self)
        win.grab_set()
        result = {"value": None}
        tk.Label(win, text=prompt, bg=COLORS["surface"], fg=COLORS["text"],
                 font=("Segoe UI", 10)).pack(padx=12, pady=(12, 4))
        var = tk.StringVar()
        entry = tk.Entry(win, textvariable=var, width=40, font=("Segoe UI", 12),
                         bg=COLORS["surface2"], fg=COLORS["text"],
                         insertbackground=COLORS["text"], relief=tk.FLAT)
        entry.pack(padx=12, pady=6)
        entry.focus_set()

        def insert_char(ch):
            if ch == "BACK":
                cur = var.get()
                var.set(cur[:-1])
            else:
                entry.insert(tk.INSERT, ch)
            entry.focus_set()

        def ok():
            result["value"] = var.get()
            win.destroy()

        def cancel():
            win.destroy()

        bf = tk.Frame(win, bg=COLORS["surface"])
        bf.pack(pady=8)
        tk.Button(bf, text="বাংলা কিবোর্ড", command=lambda: BanglaKeyboard(win, on_char=insert_char, on_enter=ok),
                  bg="#0ea5e9", fg="white", relief=tk.FLAT, padx=10, pady=4,
                  font=("Segoe UI", 9), cursor="hand2").pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="OK", command=ok, bg=COLORS["accent"], fg="white",
                  relief=tk.FLAT, padx=14, pady=4, font=("Segoe UI", 9), cursor="hand2").pack(side=tk.LEFT, padx=4)
        tk.Button(bf, text="Cancel", command=cancel, bg=COLORS["surface2"], fg=COLORS["text"],
                  relief=tk.FLAT, padx=10, pady=4, font=("Segoe UI", 9), cursor="hand2").pack(side=tk.LEFT, padx=4)
        entry.bind("<Return>", lambda e: ok())
        win.wait_window()
        return result["value"]

    def _open_bangla_kb_standalone(self):
        """Bangla keyboard that fills a buffer — paste into edit fields."""
        buf = {"text": ""}

        def insert_char(ch):
            if ch == "BACK":
                buf["text"] = buf["text"][:-1]
            else:
                buf["text"] += ch
            self.status.config(text="বাংলা: " + buf["text"][-40:])

        def on_enter():
            # copy to clipboard for easy paste
            try:
                self.clipboard_clear()
                self.clipboard_append(buf["text"])
                messagebox.showinfo(
                    "Copied",
                    "বাংলা টেক্সট কপি হয়েছে।\n\nএখন টেক্সটে ডাবল-ক্লিক করে Edit → Ctrl+V পেস্ট করুন।\n\n" + buf["text"][:200],
                    parent=self,
                )
            except Exception as e:
                messagebox.showinfo("Text", buf["text"] or "(empty)", parent=self)

        BanglaKeyboard(self, on_char=insert_char, on_enter=on_enter)

    def show_about(self):
        messagebox.showinfo(
            "About",
            f"{APP_NAME}\nVersion {APP_VERSION}  •  {APP_YEAR}\n\n"
            f"Made by {APP_STUDIO}\n\n"
            "Offline PDF, Image & Document editor.\n"
            "View • Edit text in place • Sign • Screenshot\n"
            "Color Fill Box • Copy signature from PDF\n"
            "Open PDF / Images / Word → Save as PDF\n\n"
            "No internet required.",
        )

    def _on_close(self):
        if self.pdf.dirty:
            if not messagebox.askyesno("Quit", "Unsaved changes will be lost. Quit?"):
                return
        self.pdf.close()
        self.destroy()


def main():
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        pass
    app = OnePDFEditor()
    # Windows "Open with" / double-click association: path comes as argv
    if len(sys.argv) > 1:
        arg_path = sys.argv[1].strip('"')
        if os.path.isfile(arg_path):
            app.after(200, lambda p=arg_path: app._load_path(p))
    app.mainloop()


if __name__ == "__main__":
    main()
